"""
Coletor diário de notícias via Google News RSS, multi região.

Como funciona:
1. Para cada região configurada (Brasil, EUA, etc), monta URLs do RSS com
   parâmetros locais (idioma, país)
2. Para cada termo da região, busca o feed e lê os resultados
3. Resolve a URL real de cada link (segue redirect do Google)
4. Se a região exigir tradução, traduz os títulos para inglês
5. Salva tudo num Word organizado por região, depois por termo,
   no formato 'Title (Source, May 20)'

Dependências: pip install feedparser python-docx requests deep-translator

Autor: gerado com Claude
"""

import feedparser
import urllib.parse
import requests
import os
from datetime import datetime
from email.utils import parsedate_to_datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from deep_translator import GoogleTranslator


# ============================================================
# CONFIGURAÇÃO (edite aqui)
# ============================================================

# Cada região tem seus parâmetros do Google News e sua lista própria de termos.
# Os parâmetros hl, gl e ceid mudam o idioma da interface, o país e a edição.
# Para adicionar outra região (ex Reino Unido) basta replicar o bloco.
REGIOES = {
    "Brasil": {
        "hl": "pt-BR",
        "gl": "BR",
        "ceid": "BR:pt-419",
        "traduzir": True,
        "termos": [
            "regulação de IA Brasil",
            '"marco da inteligência artificial"',
            "fusão e aquisição tecnologia",
            "Banco Central juros Selic",
            "mercado de capitais Brasil",
        ],
    },
    "Estados Unidos": {
        "hl": "en-US",
        "gl": "US",
        "ceid": "US:en",
        "traduzir": False,
        "termos": [
            "AI regulation United States",
            '"AI Act" OR "AI bill"',
            "tech mergers and acquisitions",
            "Federal Reserve interest rates",
            "US stock market",
        ],
    },
}

# Janela temporal aplicada a todas as buscas. when:1h, when:6h, when:12h, when:1d, when:7d
JANELA_TEMPORAL = "when:1d"

# Pasta onde o arquivo Word vai ser salvo
PASTA_SAIDA = os.path.dirname(os.path.abspath(__file__))

# Se True, faz HEAD em cada link pra resolver a URL real do portal
RESOLVER_LINKS_REAIS = True

# Idioma de destino das traduções
IDIOMA_DESTINO = "en"

# Número de threads paralelas
MAX_THREADS = 10

USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
)


# ============================================================
# COLETA
# ============================================================

def monta_url_rss(termo: str, regiao_cfg: dict) -> str:
    """Monta URL do RSS de busca usando os parâmetros da região."""
    query_completa = f"{termo} {JANELA_TEMPORAL}".strip()
    query_codificada = urllib.parse.quote(query_completa)
    hl = regiao_cfg["hl"]
    gl = regiao_cfg["gl"]
    ceid = regiao_cfg["ceid"]
    return (
        f"https://news.google.com/rss/search?q={query_codificada}"
        f"&hl={hl}&gl={gl}&ceid={ceid}"
    )


def coleta_noticias_termo(termo: str, regiao_cfg: dict) -> list[dict]:
    """Busca o feed RSS para um termo dentro de uma região."""
    url = monta_url_rss(termo, regiao_cfg)
    feed = feedparser.parse(url, agent=USER_AGENT)

    if feed.bozo and not feed.entries:
        print(f"    Aviso: feed com problema para '{termo}'. Pulando.")
        return []

    noticias = []
    for entry in feed.entries:
        titulo = getattr(entry, "title", "") or ""
        link = getattr(entry, "link", "") or ""
        if not titulo:
            continue
        fonte = ""
        if hasattr(entry, "source") and hasattr(entry.source, "title"):
            fonte = entry.source.title
        noticias.append({
            "titulo_original": titulo,
            "titulo_en": titulo,
            "link": link,
            "fonte": fonte,
            "data_raw": getattr(entry, "published", ""),
            "data_formatada": "",
        })
    return noticias


def resolve_link_real(url_google: str) -> str:
    try:
        resposta = requests.head(
            url_google,
            allow_redirects=True,
            timeout=10,
            headers={"User-Agent": USER_AGENT},
        )
        # Alguns servidores recusam HEAD — tenta GET sem baixar o corpo
        if resposta.status_code == 405:
            resposta = requests.get(
                url_google,
                allow_redirects=True,
                timeout=10,
                headers={"User-Agent": USER_AGENT},
                stream=True,
            )
            resposta.close()
        return resposta.url
    except Exception:
        return url_google


def resolve_links_em_paralelo(noticias: list[dict]) -> list[dict]:
    if not noticias:
        return noticias
    with ThreadPoolExecutor(max_workers=MAX_THREADS) as executor:
        futuros = {
            executor.submit(resolve_link_real, n["link"]): n
            for n in noticias
            if n["link"]
        }
        for futuro in as_completed(futuros):
            noticia = futuros[futuro]
            try:
                noticia["link"] = futuro.result()
            except Exception:
                pass
    return noticias


# ============================================================
# TRADUÇÃO E DATA
# ============================================================

def traduz_texto(texto: str) -> str:
    try:
        return GoogleTranslator(source="auto", target=IDIOMA_DESTINO).translate(texto)
    except Exception as erro:
        print(f"      Falha ao traduzir, mantendo original. Motivo: {erro}")
        return texto


def traduz_em_paralelo(noticias: list[dict]) -> list[dict]:
    if not noticias:
        return noticias
    with ThreadPoolExecutor(max_workers=MAX_THREADS) as executor:
        futuros = {
            executor.submit(traduz_texto, n["titulo_original"]): n
            for n in noticias
        }
        for futuro in as_completed(futuros):
            noticia = futuros[futuro]
            try:
                noticia["titulo_en"] = futuro.result()
            except Exception:
                noticia["titulo_en"] = noticia["titulo_original"]
    return noticias


def formata_data_curta(data_raw: str) -> str:
    """Converte 'Tue, 20 May 2026 10:30:00 GMT' em 'May 20'."""
    if not data_raw:
        return ""
    try:
        dt = parsedate_to_datetime(data_raw)
        return dt.strftime("%b %d")
    except Exception:
        return ""


# ============================================================
# GERAÇÃO DO WORD
# ============================================================

def adiciona_hiperlink(paragrafo, url: str, texto: str):
    part = paragrafo.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    novo_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    cor = OxmlElement("w:color")
    cor.set(qn("w:val"), "0563C1")
    rPr.append(cor)

    sublinhado = OxmlElement("w:u")
    sublinhado.set(qn("w:val"), "single")
    rPr.append(sublinhado)

    novo_run.append(rPr)

    elemento_texto = OxmlElement("w:t")
    elemento_texto.text = texto
    novo_run.append(elemento_texto)

    hyperlink.append(novo_run)
    paragrafo._element.append(hyperlink)


def gera_arquivo_word(resultados: dict, caminho_saida: str):
    """
    Espera um dict no formato:
    {
      "Brasil": { "termo1": [noticias], "termo2": [noticias] },
      "Estados Unidos": { "termo1": [noticias] },
    }
    """
    doc = Document()

    data_hoje = datetime.now().strftime("%d/%m/%Y")
    doc.add_heading(f"Notícias do dia, {data_hoje}", level=0)

    total = sum(
        len(noticias)
        for termos_dict in resultados.values()
        for noticias in termos_dict.values()
    )
    paragrafo_resumo = doc.add_paragraph()
    run_resumo = paragrafo_resumo.add_run(
        f"Total de {total} notícias coletadas em {len(resultados)} regiões. "
        f"Títulos em inglês."
    )
    run_resumo.italic = True

    for regiao, termos_dict in resultados.items():
        doc.add_heading(regiao, level=1)

        for termo, noticias in termos_dict.items():
            doc.add_heading(termo, level=2)

            if not noticias:
                paragrafo_vazio = doc.add_paragraph()
                run_vazio = paragrafo_vazio.add_run(
                    "Nenhuma notícia encontrada para este termo."
                )
                run_vazio.italic = True
                continue

            for noticia in noticias:
                paragrafo = doc.add_paragraph(style="List Bullet")
                if noticia["link"]:
                    adiciona_hiperlink(paragrafo, noticia["link"], noticia["titulo_en"])
                else:
                    paragrafo.add_run(noticia["titulo_en"])

                partes_sufixo = []
                if noticia["fonte"]:
                    partes_sufixo.append(noticia["fonte"])
                if noticia["data_formatada"]:
                    partes_sufixo.append(noticia["data_formatada"])

                if partes_sufixo:
                    sufixo = f" ({', '.join(partes_sufixo)})"
                    paragrafo.add_run(sufixo)

    doc.save(caminho_saida)


# ============================================================
# ORQUESTRAÇÃO
# ============================================================

def main():
    inicio = datetime.now()
    print(f"[{inicio.strftime('%H:%M:%S')}] Iniciando coleta diária de notícias")

    resultados = {}

    for nome_regiao, regiao_cfg in REGIOES.items():
        print(f"\n>> Região: {nome_regiao}")
        resultados[nome_regiao] = {}

        for termo in regiao_cfg["termos"]:
            print(f"  Buscando: {termo}")
            noticias = coleta_noticias_termo(termo, regiao_cfg)

            if RESOLVER_LINKS_REAIS and noticias:
                noticias = resolve_links_em_paralelo(noticias)

            if regiao_cfg["traduzir"] and noticias:
                print(f"    Traduzindo {len(noticias)} títulos para inglês...")
                noticias = traduz_em_paralelo(noticias)

            for n in noticias:
                n["data_formatada"] = formata_data_curta(n["data_raw"])

            resultados[nome_regiao][termo] = noticias
            print(f"    {len(noticias)} resultados prontos")

    nome_arquivo = f"noticias_{datetime.now().strftime('%Y%m%d')}.docx"
    caminho = os.path.join(PASTA_SAIDA, nome_arquivo)
    gera_arquivo_word(resultados, caminho)

    duracao = (datetime.now() - inicio).total_seconds()
    print(f"\nArquivo salvo em: {caminho}")
    print(f"Tempo total: {duracao:.1f} segundos")


if __name__ == "__main__":
    main()
